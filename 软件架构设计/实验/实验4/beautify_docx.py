from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
import re

def set_run_font(run, font_name='宋体', font_size=Pt(12), bold=False, color=RGBColor(0, 0, 0)):
    """设置run的字体属性"""
    run.font.name = font_name
    run.font.size = font_size
    run.font.bold = bold
    run.font.color.rgb = color
    # 设置中文字体
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn('w:eastAsia'), font_name)

def beautify_paragraph(paragraph, is_title=False, is_heading=False, heading_level=None):
    """美化段落格式"""
    # 设置段落对齐方式
    if is_title:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif is_heading:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    # 设置段落间距
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.5
    
    # 设置段前段后间距（标题更大一些）
    if is_heading:
        paragraph.paragraph_format.space_before = Pt(12)
        paragraph.paragraph_format.space_after = Pt(6)
    
    # 遍历所有run设置字体
    for run in paragraph.runs:
        if is_title:
            set_run_font(run, font_name='黑体', font_size=Pt(18), bold=True)
        elif is_heading:
            if heading_level == 1:
                set_run_font(run, font_name='黑体', font_size=Pt(16), bold=True)
            elif heading_level == 2:
                set_run_font(run, font_name='黑体', font_size=Pt(14), bold=True)
            elif heading_level == 3:
                set_run_font(run, font_name='黑体', font_size=Pt(12), bold=True)
            else:
                set_run_font(run, font_name='黑体', font_size=Pt(12), bold=True)
        else:
            set_run_font(run, font_name='宋体', font_size=Pt(12), bold=False)

def process_document(input_path, output_path):
    doc = Document(input_path)
    
    # 设置页面边距
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(3.17)
    
    # 识别标题模式
    title_pattern = re.compile(r'^(\d+\.\s*)?.*实验.*')
    heading1_pattern = re.compile(r'^一[、\.\s].*')
    heading2_pattern = re.compile(r'^\([\d一二三四五六七八九十]+\)\s*')
    heading3_pattern = re.compile(r'^\d+[\.\s].*')
    
    is_first_paragraph = True
    
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        
        if not text:
            continue
        
        # 判断段落类型
        is_title = False
        is_heading = False
        heading_level = None
        
        if is_first_paragraph and len(text) < 50 and '实验' in text:
            is_title = True
            is_first_paragraph = False
        elif heading1_pattern.match(text) or (len(text) < 30 and text.startswith(('一、', '二、', '三、', '四、', '五、', '六、', '七、', '八、'))):
            is_heading = True
            heading_level = 1
        elif text.startswith(('1.', '2.', '3.', '4.', '5.', '6.', '7.', '8.', '9.')) and len(text.split('\n')[0]) < 50:
            is_heading = True
            heading_level = 2
        elif text.startswith('（') and '）' in text and len(text.split('）')[0]) < 20:
            is_heading = True
            heading_level = 3
        
        beautify_paragraph(paragraph, is_title=is_title, is_heading=is_heading, heading_level=heading_level)
    
    # 处理表格中的文字
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.line_spacing = 1.5
                    for run in paragraph.runs:
                        set_run_font(run, font_name='宋体', font_size=Pt(10.5))
    
    doc.save(output_path)
    print(f"文档已美化并保存至: {output_path}")

if __name__ == '__main__':
    input_file = r'D:\MyCodeResource\软件架构设计\实验\实验4\提交\2023302855 梁景毓 ex4.docx'
    output_file = r'D:\MyCodeResource\软件架构设计\实验\实验4\提交\2023302855 梁景毓 ex4_美化版.docx'
    process_document(input_file, output_file)
